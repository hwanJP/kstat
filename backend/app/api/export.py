# backend/app/api/export.py

"""
파일 내보내기 API - DOCX, HWPX 지원
HWPX: 실제 한컴오피스 파일 분석 기반 (2011 네임스페이스)
"""

import io
import os
import zipfile
import traceback
from datetime import datetime
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Literal

from ..services.session import get_session

router = APIRouter(prefix="/api/survey", tags=["export"])


class ExportRequest(BaseModel):
    format: Literal["docx", "hwpx"] = "docx"


# ============================================================================
# DOCX 생성
# ============================================================================

def create_docx(survey_content: str, title: str = "설문지") -> io.BytesIO:
    """DOCX 파일 생성"""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as e:
        print(f"[EXPORT ERROR] python-docx import 실패: {e}")
        raise HTTPException(
            status_code=500, 
            detail="python-docx 패키지가 설치되지 않았습니다. pip install python-docx"
        )
    
    try:
        doc = Document()
        
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_para = doc.add_paragraph(f"생성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()
        
        lines = survey_content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('■'):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
            elif line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                run = p.add_run(line[2:-2])
                run.bold = True
            else:
                doc.add_paragraph(line)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        print(f"[EXPORT] DOCX 생성 완료: {len(buffer.getvalue())} bytes")
        return buffer
        
    except Exception as e:
        print(f"[EXPORT ERROR] DOCX 생성 중 오류: {e}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"DOCX 생성 오류: {str(e)}")


# ============================================================================
# HWPX 생성 (실제 한컴오피스 형식 분석 기반)
# ============================================================================

def escape_xml(text: str) -> str:
    """XML 특수문자 이스케이프"""
    if text is None:
        return ""
    return (str(text)
            .replace('&', '&amp;')
            .replace('<', '&lt;')
            .replace('>', '&gt;')
            .replace('"', '&quot;')
            .replace("'", '&#39;'))


# 공통 네임스페이스 정의 (실제 한컴 파일에서 추출)
HWPX_NAMESPACES = '''xmlns:ha="http://www.hancom.co.kr/hwpml/2011/app" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph" xmlns:hp10="http://www.hancom.co.kr/hwpml/2016/paragraph" xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head" xmlns:hhs="http://www.hancom.co.kr/hwpml/2011/history" xmlns:hm="http://www.hancom.co.kr/hwpml/2011/master-page" xmlns:hpf="http://www.hancom.co.kr/schema/2011/hpf" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:opf="http://www.idpf.org/2007/opf/" xmlns:ooxmlchart="http://www.hancom.co.kr/hwpml/2016/ooxmlchart" xmlns:hwpunitchar="http://www.hancom.co.kr/hwpml/2016/HwpUnitChar" xmlns:epub="http://www.idpf.org/2007/ops" xmlns:config="urn:oasis:names:tc:opendocument:xmlns:config:1.0"'''


def create_hwpx(survey_content: str, title: str = "설문지") -> io.BytesIO:
    """HWPX 파일 생성 (실제 한컴오피스 형식)"""
    
    try:
        now = datetime.now()
        created_date = now.strftime("%Y-%m-%dT%H:%M:%SZ")
        created_date_kr = now.strftime("%Y년 %m월 %d일 %H:%M:%S")
        
        # =====================================================================
        # 1. mimetype
        # =====================================================================
        mimetype = "application/hwp+zip"
        
        # =====================================================================
        # 2. version.xml (실제 한컴 형식)
        # =====================================================================
        version_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes" ?><hv:HCFVersion xmlns:hv="http://www.hancom.co.kr/hwpml/2011/version" tagetApplication="WORDPROCESSOR" major="5" minor="1" micro="0" buildNumber="1" os="1" xmlVersion="1.4" application="Survey Builder AI" appVersion="1.0.0"/>'''

        # =====================================================================
        # 3. META-INF/manifest.xml (거의 비어있음)
        # =====================================================================
        manifest_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes" ?><odf:manifest xmlns:odf="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0"/>'''

        # =====================================================================
        # 4. META-INF/container.xml
        # =====================================================================
        container_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes" ?><ocf:container xmlns:ocf="urn:oasis:names:tc:opendocument:xmlns:container" xmlns:hpf="http://www.hancom.co.kr/schema/2011/hpf"><ocf:rootfiles><ocf:rootfile full-path="Contents/content.hpf" media-type="application/hwpml-package+xml"/></ocf:rootfiles></ocf:container>'''

        # =====================================================================
        # 5. Contents/content.hpf
        # =====================================================================
        content_hpf = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes" ?><opf:package {HWPX_NAMESPACES} version="" unique-identifier="" id=""><opf:metadata><opf:title>{escape_xml(title)}</opf:title><opf:language>ko</opf:language><opf:meta name="creator" content="text">Survey Builder AI</opf:meta><opf:meta name="subject" content="text"/><opf:meta name="description" content="text">AI 생성 설문지</opf:meta><opf:meta name="lastsaveby" content="text">Survey Builder AI</opf:meta><opf:meta name="CreatedDate" content="text">{created_date}</opf:meta><opf:meta name="ModifiedDate" content="text">{created_date}</opf:meta><opf:meta name="date" content="text">{created_date_kr}</opf:meta><opf:meta name="keyword" content="text"/></opf:metadata><opf:manifest><opf:item id="header" href="Contents/header.xml" media-type="application/xml"/><opf:item id="section0" href="Contents/section0.xml" media-type="application/xml"/><opf:item id="settings" href="settings.xml" media-type="application/xml"/></opf:manifest><opf:spine><opf:itemref idref="header" linear="no"/><opf:itemref idref="section0" linear="yes"/></opf:spine></opf:package>'''

        # =====================================================================
        # 6. Contents/header.xml (최소 필수 스타일)
        # =====================================================================
        header_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes" ?><hh:head {HWPX_NAMESPACES} version="1.4" secCnt="1"><hh:beginNum page="1" footnote="1" endnote="1" pic="1" tbl="1" equation="1"/><hh:refList><hh:fontfaces itemCnt="7"><hh:fontface lang="HANGUL" fontCnt="1"><hh:font id="0" face="함초롬돋움" type="TTF" isEmbedded="0"><hh:typeInfo familyType="FCAT_GOTHIC" weight="6" proportion="4" contrast="0" strokeVariation="1" armStyle="1" letterform="1" midline="1" xHeight="1"/></hh:font></hh:fontface><hh:fontface lang="LATIN" fontCnt="1"><hh:font id="0" face="함초롬돋움" type="TTF" isEmbedded="0"><hh:typeInfo familyType="FCAT_GOTHIC" weight="6" proportion="4" contrast="0" strokeVariation="1" armStyle="1" letterform="1" midline="1" xHeight="1"/></hh:font></hh:fontface><hh:fontface lang="HANJA" fontCnt="1"><hh:font id="0" face="함초롬돋움" type="TTF" isEmbedded="0"><hh:typeInfo familyType="FCAT_GOTHIC" weight="6" proportion="4" contrast="0" strokeVariation="1" armStyle="1" letterform="1" midline="1" xHeight="1"/></hh:font></hh:fontface><hh:fontface lang="JAPANESE" fontCnt="1"><hh:font id="0" face="함초롬돋움" type="TTF" isEmbedded="0"><hh:typeInfo familyType="FCAT_GOTHIC" weight="6" proportion="4" contrast="0" strokeVariation="1" armStyle="1" letterform="1" midline="1" xHeight="1"/></hh:font></hh:fontface><hh:fontface lang="OTHER" fontCnt="1"><hh:font id="0" face="함초롬돋움" type="TTF" isEmbedded="0"><hh:typeInfo familyType="FCAT_GOTHIC" weight="6" proportion="4" contrast="0" strokeVariation="1" armStyle="1" letterform="1" midline="1" xHeight="1"/></hh:font></hh:fontface><hh:fontface lang="SYMBOL" fontCnt="1"><hh:font id="0" face="함초롬돋움" type="TTF" isEmbedded="0"><hh:typeInfo familyType="FCAT_GOTHIC" weight="6" proportion="4" contrast="0" strokeVariation="1" armStyle="1" letterform="1" midline="1" xHeight="1"/></hh:font></hh:fontface><hh:fontface lang="USER" fontCnt="1"><hh:font id="0" face="함초롬돋움" type="TTF" isEmbedded="0"><hh:typeInfo familyType="FCAT_GOTHIC" weight="6" proportion="4" contrast="0" strokeVariation="1" armStyle="1" letterform="1" midline="1" xHeight="1"/></hh:font></hh:fontface></hh:fontfaces><hh:borderFills itemCnt="2"><hh:borderFill id="1" threeD="0" shadow="0" centerLine="NONE" breakCellSeparateLine="0"><hh:slash type="NONE" Crooked="0" isCounter="0"/><hh:backSlash type="NONE" Crooked="0" isCounter="0"/><hh:leftBorder type="NONE" width="0.1 mm" color="#000000"/><hh:rightBorder type="NONE" width="0.1 mm" color="#000000"/><hh:topBorder type="NONE" width="0.1 mm" color="#000000"/><hh:bottomBorder type="NONE" width="0.1 mm" color="#000000"/><hh:diagonal type="SOLID" width="0.1 mm" color="#000000"/></hh:borderFill><hh:borderFill id="2" threeD="0" shadow="0" centerLine="NONE" breakCellSeparateLine="0"><hh:slash type="NONE" Crooked="0" isCounter="0"/><hh:backSlash type="NONE" Crooked="0" isCounter="0"/><hh:leftBorder type="NONE" width="0.1 mm" color="#000000"/><hh:rightBorder type="NONE" width="0.1 mm" color="#000000"/><hh:topBorder type="NONE" width="0.1 mm" color="#000000"/><hh:bottomBorder type="NONE" width="0.1 mm" color="#000000"/><hh:diagonal type="SOLID" width="0.1 mm" color="#000000"/><hc:fillBrush><hc:winBrush faceColor="none" hatchColor="#999999" alpha="0"/></hc:fillBrush></hh:borderFill></hh:borderFills><hh:charProperties itemCnt="3"><hh:charPr id="0" height="1000" textColor="#000000" shadeColor="none" useFontSpace="0" useKerning="0" symMark="NONE" borderFillIDRef="2"><hh:fontRef hangul="0" latin="0" hanja="0" japanese="0" other="0" symbol="0" user="0"/><hh:ratio hangul="100" latin="100" hanja="100" japanese="100" other="100" symbol="100" user="100"/><hh:spacing hangul="0" latin="0" hanja="0" japanese="0" other="0" symbol="0" user="0"/><hh:relSz hangul="100" latin="100" hanja="100" japanese="100" other="100" symbol="100" user="100"/><hh:offset hangul="0" latin="0" hanja="0" japanese="0" other="0" symbol="0" user="0"/><hh:underline type="NONE" shape="SOLID" color="#000000"/><hh:strikeout shape="NONE" color="#000000"/><hh:outline type="NONE"/><hh:shadow type="NONE" color="#B2B2B2" offsetX="10" offsetY="10"/></hh:charPr><hh:charPr id="1" height="1400" textColor="#000000" shadeColor="none" useFontSpace="0" useKerning="0" symMark="NONE" borderFillIDRef="2" bold="1"><hh:fontRef hangul="0" latin="0" hanja="0" japanese="0" other="0" symbol="0" user="0"/><hh:ratio hangul="100" latin="100" hanja="100" japanese="100" other="100" symbol="100" user="100"/><hh:spacing hangul="0" latin="0" hanja="0" japanese="0" other="0" symbol="0" user="0"/><hh:relSz hangul="100" latin="100" hanja="100" japanese="100" other="100" symbol="100" user="100"/><hh:offset hangul="0" latin="0" hanja="0" japanese="0" other="0" symbol="0" user="0"/><hh:underline type="NONE" shape="SOLID" color="#000000"/><hh:strikeout shape="NONE" color="#000000"/><hh:outline type="NONE"/><hh:shadow type="NONE" color="#B2B2B2" offsetX="10" offsetY="10"/></hh:charPr><hh:charPr id="2" height="1200" textColor="#000000" shadeColor="none" useFontSpace="0" useKerning="0" symMark="NONE" borderFillIDRef="2" bold="1"><hh:fontRef hangul="0" latin="0" hanja="0" japanese="0" other="0" symbol="0" user="0"/><hh:ratio hangul="100" latin="100" hanja="100" japanese="100" other="100" symbol="100" user="100"/><hh:spacing hangul="0" latin="0" hanja="0" japanese="0" other="0" symbol="0" user="0"/><hh:relSz hangul="100" latin="100" hanja="100" japanese="100" other="100" symbol="100" user="100"/><hh:offset hangul="0" latin="0" hanja="0" japanese="0" other="0" symbol="0" user="0"/><hh:underline type="NONE" shape="SOLID" color="#000000"/><hh:strikeout shape="NONE" color="#000000"/><hh:outline type="NONE"/><hh:shadow type="NONE" color="#B2B2B2" offsetX="10" offsetY="10"/></hh:charPr></hh:charProperties><hh:tabProperties itemCnt="1"><hh:tabPr id="0" autoTabLeft="0" autoTabRight="0"/></hh:tabProperties><hh:numberings itemCnt="1"><hh:numbering id="1" start="0"><hh:paraHead start="1" level="1" align="LEFT" useInstWidth="1" autoIndent="1" widthAdjust="0" textOffsetType="PERCENT" textOffset="50" numFormat="DIGIT" charPrIDRef="4294967295" checkable="0">^1.</hh:paraHead></hh:numbering></hh:numberings><hh:paraProperties itemCnt="2"><hh:paraPr id="0" tabPrIDRef="0" condense="0" fontLineHeight="0" snapToGrid="1" suppressLineNumbers="0" checked="0"><hh:align horizontal="JUSTIFY" vertical="BASELINE"/><hh:heading type="NONE" idRef="0" level="0"/><hh:breakSetting breakLatinWord="KEEP_WORD" breakNonLatinWord="KEEP_WORD" widowOrphan="0" keepWithNext="0" keepLines="0" pageBreakBefore="0" lineWrap="BREAK"/><hh:autoSpacing eAsianEng="0" eAsianNum="0"/><hp:switch><hp:case hp:required-namespace="http://www.hancom.co.kr/hwpml/2016/HwpUnitChar"><hh:margin><hc:intent value="0" unit="HWPUNIT"/><hc:left value="0" unit="HWPUNIT"/><hc:right value="0" unit="HWPUNIT"/><hc:prev value="0" unit="HWPUNIT"/><hc:next value="0" unit="HWPUNIT"/></hh:margin><hh:lineSpacing type="PERCENT" value="160" unit="HWPUNIT"/></hp:case><hp:default><hh:margin><hc:intent value="0" unit="HWPUNIT"/><hc:left value="0" unit="HWPUNIT"/><hc:right value="0" unit="HWPUNIT"/><hc:prev value="0" unit="HWPUNIT"/><hc:next value="0" unit="HWPUNIT"/></hh:margin><hh:lineSpacing type="PERCENT" value="160" unit="HWPUNIT"/></hp:default></hp:switch><hh:border borderFillIDRef="2" offsetLeft="0" offsetRight="0" offsetTop="0" offsetBottom="0" connect="0" ignoreMargin="0"/></hh:paraPr><hh:paraPr id="1" tabPrIDRef="0" condense="0" fontLineHeight="0" snapToGrid="1" suppressLineNumbers="0" checked="0"><hh:align horizontal="CENTER" vertical="BASELINE"/><hh:heading type="NONE" idRef="0" level="0"/><hh:breakSetting breakLatinWord="KEEP_WORD" breakNonLatinWord="KEEP_WORD" widowOrphan="0" keepWithNext="0" keepLines="0" pageBreakBefore="0" lineWrap="BREAK"/><hh:autoSpacing eAsianEng="0" eAsianNum="0"/><hp:switch><hp:case hp:required-namespace="http://www.hancom.co.kr/hwpml/2016/HwpUnitChar"><hh:margin><hc:intent value="0" unit="HWPUNIT"/><hc:left value="0" unit="HWPUNIT"/><hc:right value="0" unit="HWPUNIT"/><hc:prev value="400" unit="HWPUNIT"/><hc:next value="400" unit="HWPUNIT"/></hh:margin><hh:lineSpacing type="PERCENT" value="160" unit="HWPUNIT"/></hp:case><hp:default><hh:margin><hc:intent value="0" unit="HWPUNIT"/><hc:left value="0" unit="HWPUNIT"/><hc:right value="0" unit="HWPUNIT"/><hc:prev value="800" unit="HWPUNIT"/><hc:next value="800" unit="HWPUNIT"/></hh:margin><hh:lineSpacing type="PERCENT" value="160" unit="HWPUNIT"/></hp:default></hp:switch><hh:border borderFillIDRef="2" offsetLeft="0" offsetRight="0" offsetTop="0" offsetBottom="0" connect="0" ignoreMargin="0"/></hh:paraPr></hh:paraProperties><hh:styles itemCnt="1"><hh:style id="0" type="PARA" name="바탕글" engName="Normal" paraPrIDRef="0" charPrIDRef="0" nextStyleIDRef="0" langID="1042" lockForm="0"/></hh:styles></hh:refList><hh:compatibleDocument targetProgram="HWP201X"><hh:layoutCompatibility/></hh:compatibleDocument><hh:docOption><hh:linkinfo path="" pageInherit="0" footnoteInherit="0"/></hh:docOption><hh:trackchageConfig flags="56"/></hh:head>'''

        # =====================================================================
        # 7. Contents/section0.xml (본문)
        # =====================================================================
        def make_paragraph(text: str, char_pr: int = 0, para_pr: int = 0, vert_pos: int = 0) -> str:
            """단일 문단 생성"""
            text = escape_xml(text)
            text_len = len(text) * 1000  # 대략적인 텍스트 너비
            return f'<hp:p id="0" paraPrIDRef="{para_pr}" styleIDRef="0" pageBreak="0" columnBreak="0" merged="0"><hp:run charPrIDRef="{char_pr}"><hp:t>{text}</hp:t></hp:run><hp:linesegarray><hp:lineseg textpos="0" vertpos="{vert_pos}" vertsize="1000" textheight="1000" baseline="850" spacing="600" horzpos="0" horzsize="{text_len}" flags="393216"/></hp:linesegarray></hp:p>'
        
        # 문단들 생성
        paragraphs = []
        vert_pos = 0
        line_height = 1600  # 줄 높이
        
        # 제목
        paragraphs.append(make_paragraph(title, char_pr=1, para_pr=1, vert_pos=vert_pos))
        vert_pos += line_height
        
        # 생성일
        created_str = now.strftime("%Y-%m-%d %H:%M")
        paragraphs.append(make_paragraph(f"생성일: {created_str}", char_pr=0, para_pr=0, vert_pos=vert_pos))
        vert_pos += line_height
        
        # 빈 줄
        paragraphs.append(make_paragraph("", char_pr=0, para_pr=0, vert_pos=vert_pos))
        vert_pos += line_height
        
        # 본문 내용
        lines = survey_content.split('\n')
        for line in lines:
            line = line.strip()
            
            if not line:
                paragraphs.append(make_paragraph("", char_pr=0, para_pr=0, vert_pos=vert_pos))
            elif line.startswith('# '):
                paragraphs.append(make_paragraph(line[2:], char_pr=1, para_pr=1, vert_pos=vert_pos))
            elif line.startswith('## '):
                paragraphs.append(make_paragraph(line[3:], char_pr=2, para_pr=0, vert_pos=vert_pos))
            elif line.startswith('### '):
                paragraphs.append(make_paragraph(line[4:], char_pr=2, para_pr=0, vert_pos=vert_pos))
            elif line.startswith('■'):
                paragraphs.append(make_paragraph(line, char_pr=2, para_pr=0, vert_pos=vert_pos))
            elif line.startswith('**') and line.endswith('**'):
                paragraphs.append(make_paragraph(line[2:-2], char_pr=2, para_pr=0, vert_pos=vert_pos))
            else:
                paragraphs.append(make_paragraph(line, char_pr=0, para_pr=0, vert_pos=vert_pos))
            
            vert_pos += line_height
        
        paragraphs_xml = ''.join(paragraphs)
        
        # 첫 번째 문단에 secPr 추가 (페이지 설정)
        sec_pr = '''<hp:secPr id="" textDirection="HORIZONTAL" spaceColumns="1134" tabStop="8000" tabStopVal="4000" tabStopUnit="HWPUNIT" outlineShapeIDRef="1" memoShapeIDRef="0" textVerticalWidthHead="0" masterPageCnt="0"><hp:grid lineGrid="0" charGrid="0" wonggojiFormat="0"/><hp:startNum pageStartsOn="BOTH" page="0" pic="0" tbl="0" equation="0"/><hp:visibility hideFirstHeader="0" hideFirstFooter="0" hideFirstMasterPage="0" border="SHOW_ALL" fill="SHOW_ALL" hideFirstPageNum="0" hideFirstEmptyLine="0" showLineNumber="0"/><hp:lineNumberShape restartType="0" countBy="0" distance="0" startNumber="0"/><hp:pagePr landscape="NARROWLY" width="59528" height="84186" gutterType="LEFT_ONLY"><hp:margin header="4252" footer="4252" gutter="0" left="8504" right="8504" top="5668" bottom="4252"/></hp:pagePr><hp:footNotePr><hp:autoNumFormat type="DIGIT" userChar="" prefixChar="" suffixChar=")" supscript="0"/><hp:noteLine length="-1" type="SOLID" width="0.12 mm" color="#000000"/><hp:noteSpacing betweenNotes="283" belowLine="567" aboveLine="850"/><hp:numbering type="CONTINUOUS" newNum="1"/><hp:placement place="EACH_COLUMN" beneathText="0"/></hp:footNotePr><hp:endNotePr><hp:autoNumFormat type="DIGIT" userChar="" prefixChar="" suffixChar=")" supscript="0"/><hp:noteLine length="14692344" type="SOLID" width="0.12 mm" color="#000000"/><hp:noteSpacing betweenNotes="0" belowLine="567" aboveLine="850"/><hp:numbering type="CONTINUOUS" newNum="1"/><hp:placement place="END_OF_DOCUMENT" beneathText="0"/></hp:endNotePr><hp:pageBorderFill type="BOTH" borderFillIDRef="1" textBorder="PAPER" headerInside="0" footerInside="0" fillArea="PAPER"><hp:offset left="1417" right="1417" top="1417" bottom="1417"/></hp:pageBorderFill></hp:secPr>'''
        
        # 첫 문단에 secPr 삽입
        first_para = f'<hp:p id="0" paraPrIDRef="1" styleIDRef="0" pageBreak="0" columnBreak="0" merged="0"><hp:run charPrIDRef="1">{sec_pr}<hp:t>{escape_xml(title)}</hp:t></hp:run><hp:linesegarray><hp:lineseg textpos="0" vertpos="0" vertsize="1400" textheight="1400" baseline="1190" spacing="840" horzpos="0" horzsize="42520" flags="393216"/></hp:linesegarray></hp:p>'
        
        # 나머지 문단들
        rest_paragraphs = ''.join(paragraphs[1:])
        
        section0_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes" ?><hs:sec {HWPX_NAMESPACES}>{first_para}{rest_paragraphs}</hs:sec>'''

        # =====================================================================
        # 8. settings.xml
        # =====================================================================
        settings_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes" ?><ha:HWPApplicationSetting xmlns:ha="http://www.hancom.co.kr/hwpml/2011/app" xmlns:config="urn:oasis:names:tc:opendocument:xmlns:config:1.0"><ha:CaretPosition listIDRef="0" paraIDRef="0" pos="0"/></ha:HWPApplicationSetting>'''

        # =====================================================================
        # HWPX 파일 생성 (ZIP)
        # =====================================================================
        buffer = io.BytesIO()
        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as hwpx:
            # mimetype은 압축하지 않고 첫 번째로 (필수!)
            hwpx.writestr('mimetype', mimetype, compress_type=zipfile.ZIP_STORED)
            hwpx.writestr('version.xml', version_xml)
            hwpx.writestr('META-INF/manifest.xml', manifest_xml)
            hwpx.writestr('META-INF/container.xml', container_xml)
            hwpx.writestr('Contents/content.hpf', content_hpf)
            hwpx.writestr('Contents/header.xml', header_xml)
            hwpx.writestr('Contents/section0.xml', section0_xml)
            hwpx.writestr('settings.xml', settings_xml)
        
        buffer.seek(0)
        print(f"[EXPORT] HWPX 생성 완료: {len(buffer.getvalue())} bytes")
        return buffer
        
    except Exception as e:
        print(f"[EXPORT ERROR] HWPX 생성 중 오류: {e}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"HWPX 생성 오류: {str(e)}")


# ============================================================================
# 엔드포인트
# ============================================================================

@router.post("/export/{session_id}")
async def export_survey(session_id: str, request: ExportRequest):
    """설문지 내보내기 (DOCX/HWPX)"""
    print(f"[EXPORT] 요청: session_id={session_id}, format={request.format}")
    
    try:
        session = get_session(session_id)
        
        if not session:
            raise HTTPException(status_code=404, detail="세션을 찾을 수 없습니다")
        
        state = session.get("state", {})
        survey_content = state.get("final_survey") or state.get("survey_draft")
        
        if not survey_content:
            raise HTTPException(status_code=400, detail="내보낼 설문지 내용이 없습니다")
        
        print(f"[EXPORT] 설문 내용 길이: {len(survey_content)} chars")
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        if request.format == "docx":
            buffer = create_docx(survey_content, "설문지")
            filename = f"survey_{timestamp}.docx"
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            
        elif request.format == "hwpx":
            buffer = create_hwpx(survey_content, "설문지")
            filename = f"survey_{timestamp}.hwpx"
            media_type = "application/hwp+zip"
            
        else:
            raise HTTPException(status_code=400, detail="지원하지 않는 형식입니다")
        
        print(f"[EXPORT] 파일 생성 완료: {filename}")
        
        return StreamingResponse(
            buffer,
            media_type=media_type,
            headers={
                "Content-Disposition": f"attachment; filename={filename}",
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"[EXPORT ERROR] 오류: {e}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"내보내기 오류: {str(e)}")